import logging
import os
import io
import asyncio
import json
import re
from datetime import datetime
from typing import Dict, Any, List, Optional, Union

# SQLAlchemy
from sqlalchemy.future import select
from sqlalchemy.orm import selectinload
from sqlalchemy import func, desc

# App modules
from app.core.database import AsyncSessionLocal
from app.models.sql_models import Task, AnalysisResult, Post, Comment, Sentiment
from app.qa.llm_service import LLMQuestionAnswering, QAService

# ReportLab (PDF Generation)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Python-docx (Word Generation)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

class ReportService:
    def __init__(self):
        self.db_session = None
        # 获取 LLM 客户端配置
        api_key = os.environ.get('OPENAI_API_KEY')
        api_base = os.environ.get('OPENAI_API_BASE')
        # 复用 QAService 来管理 LLM
        self.qa_service = QAService(api_key=api_key, api_base=api_base)

    @staticmethod
    def _is_detailed_report(content: Optional[str]) -> bool:
        cleaned = ReportService._sanitize_markdown_content(content)
        if not cleaned:
            return False

        # 聚合分析阶段会写入一条简要摘要，这种内容不应被识别为详细报告缓存。
        if cleaned.startswith("本次分析共处理") and "\n" not in cleaned:
            return False

        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]
        heading_count = sum(
            1
            for line in lines
            if line.startswith("#")
            or re.match(r"^(?:\d+[\.\、]|[一二三四五六七八九十]+[\.\、])\s*", line)
        )
        bullet_count = sum(1 for line in lines if line.startswith(("- ", "* ", "• ")))
        markers = [
            "总体舆情态势",
            "核心观点摘要",
            "争议与风险分析",
            "建议与对策",
            "舆情综述",
            "情感深度分析",
            "核心观点提取",
            "风险预警",
            "智能解读",
        ]
        return (
            (heading_count >= 3 and len(cleaned) >= 120)
            or (heading_count >= 2 and bullet_count >= 2 and len(cleaned) >= 100)
            or (len(cleaned) >= 120 and sum(marker in cleaned for marker in markers) >= 2)
        )

    @staticmethod
    def _sanitize_markdown_content(content: Optional[str]) -> str:
        if not content:
            return ""

        text = content.replace("\r\n", "\n").replace("\r", "\n").strip()
        if not text:
            return ""

        cleaned_lines: List[str] = []
        for raw_line in text.split("\n"):
            line = raw_line.strip()

            # 丢弃围栏代码块标记，避免页面与导出文档出现 ```markdown 原样输出。
            if re.fullmatch(r"```[\w-]*", line) or line == "```":
                continue

            line = re.sub(r"^`{1,3}markdown\s*", "", line, flags=re.IGNORECASE)
            line = re.sub(r"^`{1,3}", "", line)
            line = re.sub(r"`{1,3}$", "", line).strip()

            # 统一常见的编号小节格式，便于前端和导出逻辑识别为标题。
            numbered_heading = re.match(r"^(\d+)[\.\、]\s*(.+)$", line)
            if numbered_heading:
                line = f"## {numbered_heading.group(1)}. {numbered_heading.group(2).strip()}"

            if cleaned_lines and not line and not cleaned_lines[-1]:
                continue
            cleaned_lines.append(line)

        cleaned = "\n".join(cleaned_lines).strip()
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        return cleaned

    @staticmethod
    def _json_data(value: Any, default: Any) -> Any:
        if value is None:
            return default
        if isinstance(value, str):
            try:
                return json.loads(value)
            except Exception:
                return default
        return value

    def _extract_keywords(self, analysis_result: Optional[AnalysisResult], limit: int = 15) -> List[str]:
        if not analysis_result or not analysis_result.word_cloud:
            return []
        wc_data = self._json_data(analysis_result.word_cloud, [])
        if isinstance(wc_data, dict):
            wc_data = wc_data.get("all", [])
        if not isinstance(wc_data, list):
            return []
        keywords = []
        for item in wc_data[:limit]:
            if isinstance(item, dict) and item.get("name"):
                keywords.append(str(item["name"]))
        return keywords

    def _build_fallback_comprehensive_report(
        self,
        task: Task,
        sent_dist: Dict[str, int],
        total_comments: int,
        keywords: List[str],
        top_posts: List[Post],
        top_pos_comments: List[str],
        top_neg_comments: List[str],
    ) -> str:
        dominant_sentiment = "中性"
        if sent_dist:
            label = max(sent_dist, key=sent_dist.get)
            dominant_sentiment = {"positive": "正面", "neutral": "中性", "negative": "负面"}.get(label, "中性")

        post_lines = []
        for post in top_posts:
            content = (post.content or "").replace("\n", " ").strip()
            if content:
                post_lines.append(f"- {content[:100]}")
        if not post_lines:
            post_lines.append("- 暂无代表性帖子数据")

        pos_lines = top_pos_comments or ["- 暂无高热度正面评论"]
        neg_lines = top_neg_comments or ["- 暂无高热度负面评论"]
        keyword_text = "、".join(keywords[:10]) if keywords else "暂无明显高频词"

        return (
            f"# 1. 舆情综述\n"
            f"本次任务围绕“{task.keyword}”在 {task.platform} 平台展开，共纳入 {total_comments} 条评论样本。"
            f"当前整体舆情以{dominant_sentiment}为主，核心讨论集中在 {keyword_text} 等话题。\n\n"
            f"# 2. 情感深度分析\n"
            f"正面评论 {sent_dist.get('positive', 0)} 条，中性评论 {sent_dist.get('neutral', 0)} 条，"
            f"负面评论 {sent_dist.get('negative', 0)} 条。情绪分布反映出用户对事件关注度较高，"
            f"但不同观点之间仍存在明显分化，需要结合具体反馈持续跟踪。\n\n"
            f"# 3. 核心观点提取\n"
            f"{chr(10).join(post_lines)}\n\n"
            f"正向反馈示例：\n{chr(10).join(pos_lines)}\n\n"
            f"负向反馈示例：\n{chr(10).join(neg_lines)}\n\n"
            f"# 4. 风险预警\n"
            f"若负面评论继续增长，应重点关注高频负向关键词和重复投诉点，及时澄清事实并回应用户关切。\n\n"
            f"# 5. 建议与对策\n"
            f"- 持续监测高频关键词与评论情绪变化\n"
            f"- 对集中投诉点给出明确回应口径\n"
            f"- 对优质正面反馈进行内容放大和沉淀\n"
            f"- 对争议议题建立专项复盘与预警机制"
        )

    async def get_cached_report(self, task_id: int) -> Optional[str]:
        """Read an existing detailed report without triggering regeneration."""
        async with AsyncSessionLocal() as session:
            result = await session.execute(
                select(Task).options(selectinload(Task.analysis_result)).filter(Task.id == task_id)
            )
            task = result.scalars().first()
            if not task:
                raise ValueError(f"Task with id {task_id} not found")

            if not task.analysis_result or not task.analysis_result.summary:
                return None

            cleaned = self._sanitize_markdown_content(task.analysis_result.summary)
            if not self._is_detailed_report(cleaned):
                return None

            # 顺手清理历史脏数据，避免后续展示时反复兼容处理。
            if cleaned != task.analysis_result.summary:
                task.analysis_result.summary = cleaned
                task.analysis_result.generated_at = task.analysis_result.generated_at or datetime.now()
                session.add(task.analysis_result)
                await session.commit()

            return cleaned

    async def generate_comprehensive_report(self, task_id: int, force_regenerate: bool = False) -> str:
        """
        Generate a comprehensive AI analysis report for a specific task.
        Aggregates posts, comments, sentiments, and keywords to prompt the LLM.
        Returns the generated markdown report and saves it to AnalysisResult.summary.
        """
        async with AsyncSessionLocal() as session:
            # 1. Fetch Task and related AnalysisResult
            result = await session.execute(
                select(Task).options(selectinload(Task.analysis_result)).filter(Task.id == task_id)
            )
            task = result.scalars().first()
            if not task:
                raise ValueError(f"Task with id {task_id} not found")

            # 任务执行阶段会先写入简要摘要，
            # 这里只复用已经生成完成的详细报告，避免重复调用大模型。
            if (
                not force_regenerate
                and task.analysis_result
                and self._is_detailed_report(task.analysis_result.summary)
            ):
                return task.analysis_result.summary

            # 2. Fetch Posts (Top 5 by likes/comments)
            posts_result = await session.execute(
                select(Post).filter(Post.task_id == task_id).order_by(desc(Post.likes + Post.comments_count)).limit(5)
            )
            top_posts = posts_result.scalars().all()
            
            # 3. Fetch Comments stats and examples
            # Sentiment counts
            sentiment_counts = await session.execute(
                select(Sentiment.sentiment_label, func.count(Sentiment.id))
                .join(Comment)
                .join(Post)
                .filter(Post.task_id == task_id)
                .group_by(Sentiment.sentiment_label)
            )
            sent_dist = {row[0]: row[1] for row in sentiment_counts.all()}
            total_comments = sum(sent_dist.values())
            
            # Top positive comments
            pos_comments_res = await session.execute(
                select(Comment.content, Comment.likes)
                .join(Sentiment)
                .join(Post)
                .filter(Post.task_id == task_id, Sentiment.sentiment_label == 'positive')
                .order_by(desc(Comment.likes))
                .limit(5)
            )
            top_pos_comments = [f"- {row[0]} (点赞: {row[1]})" for row in pos_comments_res.all()]
            
            # Top negative comments
            neg_comments_res = await session.execute(
                select(Comment.content, Comment.likes)
                .join(Sentiment)
                .join(Post)
                .filter(Post.task_id == task_id, Sentiment.sentiment_label == 'negative')
                .order_by(desc(Comment.likes))
                .limit(5)
            )
            top_neg_comments = [f"- {row[0]} (点赞: {row[1]})" for row in neg_comments_res.all()]

            # 4. Get Keywords from AnalysisResult
            keywords = []
            if task.analysis_result:
                keywords = self._extract_keywords(task.analysis_result, limit=15)

            # 将统计结果、代表性帖子与评论示例组织成提示词，
            # 引导大模型生成结构化分析报告。
            # 5. Construct Prompt
            prompt = f"""
你是一位资深的舆情分析师。请根据以下数据，对关于"{task.keyword}"的舆情进行深度分析。

【基本数据】
- 平台：{task.platform}
- 关键词：{task.keyword}
- 情感分布：正面 {sent_dist.get('positive', 0)}, 中性 {sent_dist.get('neutral', 0)}, 负面 {sent_dist.get('negative', 0)}
- 总评论样本数：{total_comments}

【高频关键词】
{', '.join(keywords)}

【热门帖子摘要】
{chr(10).join([f"- {p.content[:100]}..." for p in top_posts])}

【精选正面评论】
{chr(10).join(top_pos_comments)}

【精选负面评论】
{chr(10).join(top_neg_comments)}

请生成一份结构清晰的Markdown格式报告，包含以下部分：
## 1. 总体舆情态势
(描述整体情感倾向，主要讨论点)

## 2. 核心观点摘要
(总结用户的主要观点，包括赞扬点和吐槽点)

## 3. 争议与风险分析
(如果存在大量负面评论，分析原因；如果存在激烈争论，分析焦点)

## 4. 建议与对策
(针对负面舆情或用户痛点，提出具体的改进或公关建议)
"""
            # 优先调用大模型生成自然语言报告；
            # 如果外部模型不可用，则退化为基于规则的保底报告。
            # 6. Call LLM
            try:
                # Use ask_custom_question with use_context=False
                llm_response = await self.qa_service.ask_custom_question(prompt, use_context=False)
                report_content = llm_response.get('answer', '生成报告失败')
            except Exception as e:
                logger.error(f"Failed to generate comprehensive report via LLM: {e}")
                report_content = self._build_fallback_comprehensive_report(
                    task=task,
                    sent_dist=sent_dist,
                    total_comments=total_comments,
                    keywords=keywords,
                    top_posts=top_posts,
                    top_pos_comments=top_pos_comments,
                    top_neg_comments=top_neg_comments,
                )

            report_content = self._sanitize_markdown_content(report_content)
            if not report_content:
                report_content = self._build_fallback_comprehensive_report(
                    task=task,
                    sent_dist=sent_dist,
                    total_comments=total_comments,
                    keywords=keywords,
                    top_posts=top_posts,
                    top_pos_comments=top_pos_comments,
                    top_neg_comments=top_neg_comments,
                )

            if not task.analysis_result:
                task.analysis_result = AnalysisResult(task_id=task_id)

            task.analysis_result.summary = report_content
            task.analysis_result.generated_at = datetime.now()

            session.add(task.analysis_result)
            await session.commit()

            return report_content

    async def generate_hot_topic_analysis(self, title: str, summary: str = None, extra_data: Dict = None) -> Dict[str, str]:
        """
        Generate analysis for a hot topic using LLM.
        """
        url = extra_data.get('url', '') if extra_data else ''
        
        prompt = f"请对以下热点话题进行深度分析：\n标题：{title}\n"
        if url:
            prompt += f"链接：{url}\n"
        if summary:
            prompt += f"摘要/简介：{summary}\n"
        if extra_data:
            # Avoid dumping too much raw data
            clean_extra = {k: v for k, v in extra_data.items() if k in ['hover', 'desc', 'description'] and v != summary}
            if clean_extra:
                prompt += f"额外信息：{json.dumps(clean_extra, ensure_ascii=False)}\n"
            
        prompt += "\n请提供以下两部分内容：\n1. 内容提炼：自动提取关键信息，用简洁的语言概括核心事件。\n2. 智能解读：生成深度分析报告，包括事件背景、可能的影响、公众情绪倾向预测等。"
        
        try:
            # use_context=False because we are providing all context in the prompt itself
            # and we don't want the RAG system to look for (non-existent) documents
            result = await self.qa_service.ask_custom_question(prompt, use_context=False)
            answer = result.get('answer', '')
            
            content_summary = ""
            analysis_report = ""
            
            if "1. 内容提炼" in answer and "2. 智能解读" in answer:
                parts = answer.split("2. 智能解读")
                content_summary = parts[0].replace("1. 内容提炼", "").strip("：:\n ")
                analysis_report = parts[1].strip("：:\n ")
            elif "内容提炼" in answer and "智能解读" in answer:
                # Fallback for looser matching
                parts = answer.split("智能解读")
                content_summary = parts[0].replace("内容提炼", "").strip("：:\n ")
                analysis_report = parts[1].strip("：:\n ")
            else:
                content_summary = answer[:200] + "..." # Fallback
                analysis_report = answer

            content_summary = self._sanitize_markdown_content(content_summary)
            analysis_report = self._sanitize_markdown_content(analysis_report)
                
            return {
                "summary": content_summary,
                "analysis": analysis_report
            }
        except Exception as e:
            logger.error(f"Hot topic analysis failed: {e}")
            return {
                "summary": "分析失败",
                "analysis": str(e)
            }

    def export_hot_topic_report(self, title: str, summary: str, analysis: str, format: str = 'pdf') -> bytes:
        if format == 'pdf':
            return self._create_hot_topic_pdf(title, summary, analysis)
        elif format == 'word':
            if Document is None:
                raise ImportError("python-docx is not installed")
            return self._create_hot_topic_word(title, summary, analysis)
        else:
            raise ValueError("Unsupported format")
            
    def _create_hot_topic_pdf(self, title: str, summary: str, analysis: str) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72
        )
        
        font_name = self._register_chinese_font()
        styles = getSampleStyleSheet()
        styles['Normal'].fontName = font_name
        styles['Normal'].fontSize = 10
        styles['Normal'].leading = 14
        styles['Heading1'].fontName = font_name
        styles['Heading1'].fontSize = 16
        styles['Heading2'].fontName = font_name
        styles['Heading2'].fontSize = 14
        
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, alignment=1, spaceAfter=30)
        
        story = []
        story.append(Paragraph(f"{title} - 热点分析报告", title_style))
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 0.5 * inch))
        
        story.append(Paragraph("1. 内容提炼", styles['Heading1']))
        story.append(Paragraph(summary, styles['Normal']))
        story.append(Spacer(1, 0.2 * inch))
        
        story.append(Paragraph("2. 智能解读", styles['Heading1']))
        # Simple Markdown parsing
        lines = analysis.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            if line.startswith('# '): story.append(Paragraph(line[2:], styles['Heading1']))
            elif line.startswith('## '): story.append(Paragraph(line[3:], styles['Heading2']))
            elif line.startswith('- ') or line.startswith('* '): story.append(Paragraph(f"• {line[2:]}", styles['Normal']))
            else: story.append(Paragraph(line, styles['Normal']))
            
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def _create_hot_topic_word(self, title: str, summary: str, analysis: str) -> bytes:
        doc = Document()
        doc.add_heading(f"{title} - 热点分析报告", 0)
        
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        doc.add_heading('1. 内容提炼', level=1)
        doc.add_paragraph(summary)
        
        doc.add_heading('2. 智能解读', level=1)
        # Simple Markdown parsing
        lines = analysis.split('\n')
        for line in lines:
            line = line.strip()
            if not line: continue
            if line.startswith('# '): doc.add_heading(line[2:], level=1)
            elif line.startswith('## '): doc.add_heading(line[3:], level=2)
            elif line.startswith('### '): doc.add_heading(line[4:], level=3)
            elif line.startswith('- ') or line.startswith('* '): doc.add_paragraph(line[2:], style='List Bullet')
            else: doc.add_paragraph(line)
            
        buffer = io.BytesIO()
        doc.save(buffer)
        word_bytes = buffer.getvalue()
        buffer.close()
        return word_bytes

    async def generate_task_report(self, task_id: int) -> Dict[str, Any]:
        """
        生成任务分析报告的主入口
        返回: PDF 文件的二进制流 (bytes) 或 错误信息
        """
        try:
            report_markdown = await self.generate_comprehensive_report(task_id, force_regenerate=False)

            # 1. 获取完整数据
            task_data = await self._fetch_task_data(task_id)
            if not task_data:
                return {"error": "Task not found"}

            # 基于已落库的报告内容导出文档，避免导出时再次请求大模型。
            pdf_bytes = await self._create_pdf(task_data, report_markdown)
            
            return {
                "filename": f"report_{task_data['task'].keyword}_{datetime.now().strftime('%Y%m%d')}.pdf",
                "content": pdf_bytes,
                "task_id": task_id
            }
            
        except Exception as e:
            logger.error(f"Failed to generate report for task {task_id}: {e}", exc_info=True)
            return {"error": str(e)}

    async def _fetch_task_data(self, task_id: int) -> Dict[str, Any]:
        """从数据库获取任务的所有相关数据"""
        async with AsyncSessionLocal() as session:
            # 获取任务及分析结果
            stmt = select(Task).options(
                selectinload(Task.analysis_result)
            ).filter_by(id=task_id)
            result = await session.execute(stmt)
            task = result.scalars().first()

            if not task:
                return None
            
            # 获取评论数和帖子数
            post_count = await session.scalar(select(func.count(Post.id)).filter_by(task_id=task_id))
            comment_count = await session.scalar(select(func.count(Comment.id)).join(Post).filter(Post.task_id == task_id))
            
            # 获取部分代表性评论 (比如点赞最高的 5 条)
            stmt_comments = select(Comment).join(Post).filter(Post.task_id == task_id).order_by(Comment.likes.desc()).limit(5)
            comments_res = await session.execute(stmt_comments)
            top_comments = [c.content for c in comments_res.scalars().all()]

            return {
                "task": task,
                "post_count": post_count,
                "comment_count": comment_count,
                "analysis": task.analysis_result,
                "top_comments": top_comments
            }

    async def _generate_ai_content(self, data: Dict[str, Any]) -> str:
        """调用 LLM 生成深度分析报告"""
        task = data['task']
        analysis = data['analysis']
        sentiment_data = self._json_data(analysis.sentiment_distribution, {}) if analysis else {}
        word_cloud_data = self._json_data(analysis.word_cloud, []) if analysis else []
        top_words = word_cloud_data.get("all", []) if isinstance(word_cloud_data, dict) else word_cloud_data
        
        # 将任务统计结果整理为统一上下文，便于后续生成报告提示词。
        context = {
            "platform": task.platform,
            "keyword": task.keyword,
            "stats": {
                "posts": data['post_count'],
                "comments": data['comment_count']
            },
            "sentiment": sentiment_data if isinstance(sentiment_data, dict) else {},
            "top_words": top_words[:10] if isinstance(top_words, list) else [],
            "top_comments": data['top_comments']
        }
        
        prompt = f"""
        请作为一位专业的数据分析师，为以下社交媒体舆情任务撰写一份深度分析报告。
        
        【任务信息】
        - 平台: {context['platform']}
        - 关键词: {context['keyword']}
        - 数据量: {context['stats']['posts']} 帖子, {context['stats']['comments']} 评论
        
        【情感数据】
        {context['sentiment']}
        
        【高频热词】
        {context['top_words']}
        
        【代表性评论】
        {context['top_comments']}
        
        请生成一份结构化的报告（Markdown格式），包含以下部分（请严格使用以下标题）：
        
        # 1. 舆情综述
        (简要总结本次舆情的热度、总体情感倾向和主要讨论点)
        
        # 2. 情感深度分析
        (详细分析正面、负面、中性情绪的占比及其背后的原因，结合热词进行解读)
        
        # 3. 核心观点提取
        (基于代表性评论和热词，总结用户的核心诉求或槽点)
        
        # 4. 风险预警
        (如果存在负面情绪，指出潜在的风险点)
        
        # 5. 建议与对策
        (针对当前舆情，给出 3-4 条具体的运营或公关建议)
        
        注意：
        - 语言风格要专业、客观。
        - 不要包含 markdown 的代码块标记（如 ```markdown），直接返回内容。
        - 重点突出数据背后的洞察。
        """
        
        try:
            response = await self.qa_service.ask_custom_question(prompt)
            return response.get('answer', 'AI 分析生成失败，请稍后重试。')
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            return self._build_fallback_comprehensive_report(
                task=task,
                sent_dist=context["sentiment"],
                total_comments=data["comment_count"],
                keywords=[item.get("name", "") for item in context["top_words"] if isinstance(item, dict)],
                top_posts=[],
                top_pos_comments=[],
                top_neg_comments=[],
            )

    async def _create_pdf(self, data: Dict[str, Any], ai_text: str) -> bytes:
        """使用 ReportLab 生成 PDF 文件流"""
        
        # 在 Executor 中运行同步的 PDF 生成代码，避免阻塞 EventLoop
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, self._create_pdf_sync, data, ai_text)

    def _create_pdf_sync(self, data: Dict[str, Any], ai_text: str) -> bytes:
        """同步的 PDF 生成逻辑"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72, leftMargin=72,
            topMargin=72, bottomMargin=72
        )
        
        # --- 1. 注册中文字体 ---
        font_name = self._register_chinese_font()
        
        # --- 2. 定义样式 ---
        styles = getSampleStyleSheet()
        # 覆盖默认样式以使用中文字体
        styles['Normal'].fontName = font_name
        styles['Normal'].fontSize = 10
        styles['Normal'].leading = 14
        styles['Normal'].spaceAfter = 6
        
        styles['Heading1'].fontName = font_name
        styles['Heading1'].fontSize = 16
        styles['Heading1'].leading = 20
        styles['Heading1'].spaceAfter = 12
        styles['Heading1'].spaceBefore = 12
        styles['Heading1'].textColor = colors.HexColor('#2c3e50')
        
        styles['Heading2'].fontName = font_name
        styles['Heading2'].fontSize = 14
        styles['Heading2'].leading = 18
        styles['Heading2'].spaceAfter = 10
        styles['Heading2'].spaceBefore = 10
        styles['Heading2'].textColor = colors.HexColor('#34495e')
        
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=24,
            leading=30,
            alignment=1, # Center
            spaceAfter=30
        )
        
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.gray,
            alignment=1, # Center
            spaceAfter=20
        )

        # --- 3. 构建内容 ---
        story = []
        task = data['task']
        
        # 标题
        story.append(Paragraph(f"{task.keyword} - 舆情分析报告", title_style))
        story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style))
        story.append(Spacer(1, 0.5 * inch))
        
        # 基本信息表格
        table_data = [
            ['平台', task.platform, '帖子数', str(data['post_count'])],
            ['关键词', task.keyword, '评论数', str(data['comment_count'])],
            ['状态', task.status, '开始时间', task.started_at.strftime('%Y-%m-%d') if task.started_at else '-']
        ]
        t = Table(table_data, colWidths=[2*cm, 5*cm, 2*cm, 5*cm])
        t.setStyle(TableStyle([
            ('FONTNAME', (0,0), (-1,-1), font_name),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f0f2f5')),
            ('BACKGROUND', (2,0), (2,-1), colors.HexColor('#f0f2f5')),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('PADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.5 * inch))
        
        # 按照 Markdown 标题和段落结构解析报告内容，
        # 使导出的 PDF 在版式上保持清晰可读。
        lines = ai_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # 一级标题 (对应 Heading1)
                text = line[2:]
                story.append(Paragraph(text, styles['Heading1']))
            elif line.startswith('## '):
                # 二级标题 (对应 Heading2)
                text = line[3:]
                story.append(Paragraph(text, styles['Heading2']))
            elif line.startswith('- ') or line.startswith('* '):
                # 列表项
                text = f"• {line[2:]}"
                story.append(Paragraph(text, styles['Normal']))
            else:
                # 普通段落
                # 处理加粗 **text** -> <b>text</b>
                # ReportLab 支持简单的 XML 标签
                line = line.replace('**', '<b>', 1).replace('**', '</b>', 1)
                story.append(Paragraph(line, styles['Normal']))
        
        # --- 4. 生成 ---
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def _register_chinese_font(self) -> str:
        """注册中文字体，返回可用的字体名称"""
        # 尝试列表
        fonts = [
            ('SimHei', 'C:\\Windows\\Fonts\\simhei.ttf'), # Windows
            ('SimSun', 'C:\\Windows\\Fonts\\simsun.ttc'), # Windows
            ('Microsoft YaHei', 'C:\\Windows\\Fonts\\msyh.ttc'), # Windows
            ('PingFang', '/System/Library/Fonts/PingFang.ttc'), # Mac
            ('STHeiti', '/System/Library/Fonts/STHeiti Light.ttc'), # Mac Old
            ('SimHei', '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc'), # Linux Common
            ('Arial Unicode', 'Arial Unicode.ttf') # Fallback
        ]
        
        for name, path in fonts:
            if os.path.exists(path):
                try:
                    pdfmetrics.registerFont(TTFont(name, path))
                    logger.info(f"Registered font: {name} from {path}")
                    return name
                except Exception as e:
                    logger.warning(f"Failed to load font {name}: {e}")
                    continue
        
        # 如果都没找到，使用内置字体（不支持中文，会显示方框）
        logger.warning("No Chinese font found! Using default Helvetica.")
        return "Helvetica"
